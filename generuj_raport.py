#!/usr/bin/env python3
"""
Generator raportu DOCX: PV + Magazyn Energii (BESS) dla zakładów produkcyjnych w Polsce.
Raport zawiera pełny research prawny, technologiczny i rynkowy.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def add_styled_table(doc, headers, rows, col_widths=None):
    """Dodaje sformatowaną tabelę do dokumentu."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Nagłówki
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Dane
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_data)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    return table


def create_report():
    doc = Document()

    # Styl domyślny
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ===== STRONA TYTUŁOWA =====
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('PV + Magazyn Energii (BESS)\ndla Zakładów Produkcyjnych w Polsce')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Raport: Analiza prawna, technologiczna i rynkowa\nLuty 2026')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(
        'Strategia: Autokonsumpcja PV + Arbitraż cenowy z magazynem energii\n'
        'Taryfa dynamiczna (ceny 15-minutowe TGE) + Peak Shaving'
    )
    run.font.size = Pt(12)
    run.font.italic = True

    doc.add_page_break()

    # ===== SPIS TREŚCI =====
    doc.add_heading('Spis treści', level=1)
    toc_items = [
        '1. Podsumowanie wykonawcze',
        '2. Strategia działania – jak to działa?',
        '3. Regulacje prawne',
        '   3.1. Taryfy dynamiczne w Polsce',
        '   3.2. Regulacje magazynów energii',
        '   3.3. Net-billing vs autokonsumpcja',
        '   3.4. Dotacje i ulgi podatkowe',
        '   3.5. Obowiązki prawne – podsumowanie',
        '4. Analiza technologiczna',
        '   4.1. Producenci BESS dla przemysłu',
        '   4.2. Technologia LFP vs NMC',
        '   4.3. Oprogramowanie EMS',
        '   4.4. Integracja z istniejącym PV (AC vs DC coupling)',
        '   4.5. Sizing magazynu',
        '5. Analiza rynkowa',
        '   5.1. Ceny energii na TGE',
        '   5.2. Ujemne ceny energii',
        '   5.3. Arbitraż cenowy',
        '   5.4. Peak shaving i opłata mocowa',
        '   5.5. Prognozy cen 2025-2027',
        '6. Business case – strumienie przychodów',
        '7. Rekomendacje',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)

    doc.add_page_break()

    # ===== 1. PODSUMOWANIE WYKONAWCZE =====
    doc.add_heading('1. Podsumowanie wykonawcze', level=1)

    doc.add_paragraph(
        'Zakłady produkcyjne posiadające instalacje fotowoltaiczne (PV) stoją przed wyjątkową '
        'okazją inwestycyjną. Dynamicznie zmieniający się rynek energii w Polsce – z rosnącą '
        'liczbą godzin z ujemnymi cenami energii, wzrastającą opłatą mocową (+55% w 2026 r.) '
        'i dostępnością taryf dynamicznych z cenami 15-minutowymi – tworzy idealne warunki '
        'do instalacji przemysłowych magazynów energii (BESS).'
    )

    doc.add_heading('Kluczowe liczby:', level=3)
    bullets = [
        'Ujemne ceny energii: >315 godzin w 2025 r. (6-krotny wzrost vs 2023)',
        'Opłata mocowa 2026: 219,40 PLN/MWh (+55% r/r) – magazyn może ją zredukować o 83%',
        'Spread cenowy intraday: 150-600 PLN/MWh (potencjał arbitrażu)',
        'Taryfy dynamiczne: ceny co 15 min od października 2025 – 6 głównych sprzedawców',
        'Koszt BESS (C&I): 1 500 – 3 500 PLN/kWh (spadek 31% r/r)',
        'Okres zwrotu: 4-8 lat (bez dotacji), 3-5 lat (z dofinansowaniem)',
        'Technologia: LFP – 10 000 cykli, 20+ lat żywotności, najwyższe bezpieczeństwo',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_page_break()

    # ===== 2. STRATEGIA DZIAŁANIA =====
    doc.add_heading('2. Strategia działania – jak to działa?', level=1)

    doc.add_paragraph(
        'Proponowana strategia łączy trzy źródła oszczędności w jeden zintegrowany system:'
    )

    doc.add_heading('Schemat działania:', level=3)

    steps = [
        ('Godziny 10:00-15:00 (tanie/ujemne ceny)',
         'PV produkuje energię → priorytet: zasilanie zakładu (autokonsumpcja). '
         'Nadwyżki ładują magazyn BESS zamiast trafiać do sieci po niskich cenach net-billing. '
         'Dodatkowo: magazyn kupuje tanią energię z sieci po cenach 15-min z TGE (nawet po 0 zł lub ujemnych cenach).'),
        ('Godziny 17:00-21:00 (drogie godziny szczytowe)',
         'Magazyn rozładowuje się, zasilając zakład → firma nie kupuje drogiej energii ze szczytu. '
         'Redukcja opłaty mocowej (peak shaving) – wygładzenie profilu zużycia.'),
        ('Godziny nocne (22:00-06:00)',
         'Magazyn ładuje się po niskich cenach nocnych z TGE, przygotowując się na poranną zmianę.'),
    ]

    for title, desc in steps:
        p = doc.add_paragraph()
        run = p.add_run(title + ': ')
        run.bold = True
        p.add_run(desc)

    doc.add_heading('Trzy strumienie oszczędności:', level=3)
    savings = [
        'Autokonsumpcja PV: nadwyżki do magazynu zamiast do sieci (wartość 300-500 PLN/MWh wyższa niż net-billing)',
        'Arbitraż cenowy: kupno taniej energii (ceny bliskie 0 lub ujemne) i zużycie w szczycie',
        'Peak shaving: redukcja opłaty mocowej nawet o 83% (przejście z kategorii K4 na K1)',
    ]
    for s in savings:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_page_break()

    # ===== 3. REGULACJE PRAWNE =====
    doc.add_heading('3. Regulacje prawne', level=1)

    # 3.1 Taryfy dynamiczne
    doc.add_heading('3.1. Taryfy dynamiczne w Polsce', level=2)

    doc.add_paragraph(
        'Od 24 sierpnia 2024 r. w Polsce obowiązuje wymóg oferowania umów z ceną dynamiczną '
        'przez sprzedawców energii obsługujących ponad 200 tys. klientów. Od października 2025 r. '
        'ceny zmieniają się co 15 minut (96 zmian ceny dziennie), oparte na notowaniach '
        'Rynku Dnia Następnego (RDN) na Towarowej Giełdzie Energii (TGE).'
    )

    doc.add_heading('Sprzedawcy oferujący taryfy dynamiczne dla firm:', level=3)

    add_styled_table(doc,
        ['Sprzedawca', 'Oferta', 'Opłata handlowa'],
        [
            ['Enea', 'Ceny Dynamiczne Firma', '25 zł netto/mies. za układ pomiarowy'],
            ['Energa', 'Oferta Dynamiczna dla Firm', '10 zł brutto/mies.'],
            ['PGE', 'Cena dynamiczna', '50 zł brutto/mies.'],
            ['Tauron', 'Ceny Dynamiczne dla Firm', '0-34 zł/mies.'],
            ['E.ON Polska', 'Oferta dynamiczna', 'Indywidualnie'],
            ['Pstryk (Fortum)', 'Ceny dynamiczne', '0 zł/mies.'],
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        'Warunki przystąpienia: inteligentny licznik (zdalny odczyt), jednomiesięczny cykl rozliczeniowy. '
        'Taryfa dynamiczna + magazyn energii = możliwość zakupu energii po cenach bliskich 0 zł lub ujemnych.'
    )

    # 3.2 Regulacje magazynów
    doc.add_heading('3.2. Regulacje magazynów energii', level=2)

    doc.add_paragraph('Progi regulacyjne:')
    add_styled_table(doc,
        ['Moc magazynu', 'Koncesja', 'Rejestr OSD', 'Prawo budowlane'],
        [
            ['Do 50 kW', 'NIE', 'NIE', 'Do 30 kWh: brak wymagań\n30-300 kWh: zgłoszenie + PSP'],
            ['50 kW – 10 MW', 'NIE', 'TAK (wpis w 7 dni)', 'Powyżej 300 kWh: pozwolenie na budowę'],
            ['Powyżej 10 MW', 'TAK (URE)', 'TAK', 'Pozwolenie na budowę'],
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        'Nowelizacja Prawa budowlanego z 7 stycznia 2026 r. podniosła próg zwolnienia z wymagań '
        'budowlanych do 30 kWh (wcześniej 10 kWh) dla wolnostojących magazynów poza budynkami. '
        'Magazyny 30-300 kWh wymagają zgłoszenia budowy i zatwierdzenia Państwowej Straży Pożarnej.'
    )

    # 3.3 Net-billing vs autokonsumpcja
    doc.add_heading('3.3. Net-billing vs autokonsumpcja z magazynem', level=2)

    doc.add_paragraph(
        'W systemie net-billing energia oddana do sieci wyceniana jest wg Rynkowych Cen Energii (RCE) '
        'ze współczynnikiem 1,23. Wartość trafia na depozyt prosumencki (przepada po 12 miesiącach). '
        'Problem: w godzinach szczytu PV (11:00-15:00) ceny RCE są bardzo niskie, często bliskie zeru.'
    )

    doc.add_heading('Porównanie wartości energii:', level=3)
    add_styled_table(doc,
        ['Scenariusz', 'Wartość energii', 'Uwagi'],
        [
            ['Autokonsumpcja (własne zużycie)', '700-900 PLN/MWh', 'Pełna wartość: energia + dystrybucja + opłaty'],
            ['Oddanie do sieci (net-billing)', '200-400 PLN/MWh', 'Tylko składnik energetyczny wg RCE × 1,23'],
            ['Magazynowanie nadwyżek PV', '700-900 PLN/MWh', 'Taka sama wartość jak autokonsumpcja'],
        ]
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Wniosek: ')
    run.bold = True
    p.add_run(
        'Każda kWh zmagazynowana i autokonsumowana zamiast oddana do sieci = oszczędność 300-500 PLN/MWh. '
        'Magazyn energii zwiększa autokonsumpcję PV z 25-40% do 60-80%.'
    )

    # 3.4 Dotacje
    doc.add_heading('3.4. Dotacje i ulgi podatkowe', level=2)

    doc.add_heading('Dostępne programy wsparcia:', level=3)
    add_styled_table(doc,
        ['Program', 'Forma', 'Kwota', 'Dla kogo', 'Status 2026'],
        [
            ['Energia Plus (NFOŚiGW)', 'Pożyczka preferencyjna', '0,5-500 mln PLN, umorzenie do 10%', 'Przedsiębiorcy', 'Nabór zakończony, możliwe wznowienie'],
            ['Magazyny Energii (KPO)', 'Dotacja do 45-65%', 'Min. 2 MW / 4 MWh', 'Przedsiębiorcy (duże proj.)', 'Alokacja wyczerpana'],
            ['Ulga termomodernizacyjna', 'Odliczenie od dochodu', 'Do 53 000 PLN', 'Osoby fizyczne (JDG)', 'Aktywna od 01.2025'],
            ['FEnIKS / FENG / RPO', 'Dotacje UE', 'Zależne od województwa', 'MŚP i duże firmy', 'Nabory regionalne'],
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        'Uwaga: Główne programy dotacyjne na duże magazyny (KPO 4,15 mld PLN) mają wyczerpane alokacje. '
        'Dla zakładów produkcyjnych najlepsze opcje to regionalne fundusze UE (FEnIKS, RPO) '
        'oraz ewentualne wznowienie programu Energia Plus.'
    )

    # 3.5 Obowiązki prawne
    doc.add_heading('3.5. Obowiązki prawne – podsumowanie', level=2)

    doc.add_paragraph('Typowy scenariusz: zakład produkcyjny z PV (do 50 kW) + BESS (100-300 kWh):')

    add_styled_table(doc,
        ['Obowiązek', 'Wymaganie', 'Podstawa prawna'],
        [
            ['Koncesja na magazynowanie', 'NIE (moc < 10 MW)', 'Prawo energetyczne'],
            ['Wpis do rejestru OSD', 'NIE (moc ≤ 50 kW)', 'Prawo energetyczne'],
            ['Pozwolenie na budowę', 'NIE (< 30 kWh) / Zgłoszenie (30-300 kWh)', 'Prawo budowlane (01.2026)'],
            ['Warunki przyłączenia', 'Aktualizacja istniejących', 'Prawo energetyczne'],
            ['Umowa z OSD/sprzedawcą', 'Aktualizacja wymagana', 'Prawo energetyczne'],
            ['Zatwierdzenie PSP', 'TAK (dla 30-300 kWh)', 'Prawo budowlane'],
        ]
    )

    doc.add_page_break()

    # ===== 4. ANALIZA TECHNOLOGICZNA =====
    doc.add_heading('4. Analiza technologiczna', level=1)

    # 4.1 Producenci BESS
    doc.add_heading('4.1. Producenci BESS dla przemysłu', level=2)

    doc.add_paragraph(
        'Na rynku dostępnych jest kilku wiodących producentów systemów BESS w segmencie C&I '
        '(Commercial & Industrial, pojemności 100 kWh – 5 MWh):'
    )

    add_styled_table(doc,
        ['Producent', 'Model C&I', 'Pojemność', 'Cykle', 'RTE', 'Chłodzenie'],
        [
            ['CATL', 'EnerOne', '232-407 kWh', '10 000', '~95%', 'Cieczowe'],
            ['BYD', 'MC Cube-T', '5 MWh (kontener)', '6 000+', '~95%', 'Cieczowe'],
            ['Tesla', 'Megapack 3', '5 MWh', '7 000+', '~92%', 'Cieczowe'],
            ['Samsung SDI', 'SBB 2.0', '3,8-6,1 MWh', '8 000+', '>95%', 'Cieczowe EDI'],
            ['Huawei', 'LUNA2000-200', '97-200 kWh', '6 000+', '~95%', 'Cieczowe'],
            ['Sungrow', 'PowerStack 255CS', '257 kWh (do 6,4 MWh)', '6 000+', '90%', 'Cieczowe'],
            ['Pylontech', 'Force H2', 'Do 215 kWh', '6 000+', '>95%', 'Powietrzne'],
        ]
    )

    doc.add_paragraph()

    doc.add_heading('Rekomendacje dla zakładów produkcyjnych:', level=3)
    recs = [
        'CATL EnerOne (232-407 kWh) – najlepsza żywotność (10 000 cykli), dystrybutor w PL: 7SUN',
        'Sungrow PowerStack 255CS (257 kWh, skalowalne do 6,4 MWh) – świetny stosunek ceny do jakości',
        'Huawei LUNA2000-200 (97-200 kWh) – pełna integracja z ekosystemem FusionSolar',
    ]
    for r in recs:
        doc.add_paragraph(r, style='List Bullet')

    # 4.2 LFP vs NMC
    doc.add_heading('4.2. Technologia baterii: LFP vs NMC', level=2)

    add_styled_table(doc,
        ['Parametr', 'LFP (LiFePO₄)', 'NMC'],
        [
            ['Gęstość energii', '90-120 Wh/kg', '150-220 Wh/kg'],
            ['Cykle życia (do 80% SOH)', '4 000 – 10 000', '1 000 – 5 000'],
            ['Temp. thermal runaway', '270°C', '210°C'],
            ['Ryzyko thermal runaway', '80% mniejsze', 'Wyższe'],
            ['Koszt ogniw (2025)', '~60-80 USD/kWh', '~80-120 USD/kWh'],
            ['Degradacja roczna', '~1,5-2%', '~2-3%'],
            ['Kobalt', 'NIE', 'TAK'],
        ]
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Rekomendacja: LFP bezwzględnie. ')
    run.bold = True
    p.add_run(
        'Dla stacjonarnych magazynów przemysłowych LFP wygrywa we wszystkich kluczowych parametrach: '
        'bezpieczeństwo, żywotność, koszt cyklu. Gabaryty nie są ograniczeniem w zakładzie produkcyjnym.'
    )

    # 4.3 Oprogramowanie EMS
    doc.add_heading('4.3. Oprogramowanie EMS (Energy Management System)', level=2)

    doc.add_paragraph(
        'System EMS jest kluczowym elementem strategii – to oprogramowanie decyduje kiedy ładować/rozładowywać '
        'magazyn na podstawie cen z TGE, prognoz pogody, profilu zużycia zakładu i prognozy produkcji PV.'
    )

    add_styled_table(doc,
        ['System', 'Ceny TGE', 'AI/ML', 'Arbitraż', 'Dla PL'],
        [
            ['ENNO-EMS (Ennovation Tech.)', 'TAK (natywnie)', 'Tak', 'TAK', 'TAK – polski producent'],
            ['Tibo Energy Alice', 'Tak (EU)', 'TAK (najlepsza AI)', 'TAK', 'Częściowo'],
            ['gridX XENON', 'Tak (EU)', 'Tak', 'Tak', 'Częściowo'],
            ['Huawei FusionSolar', 'Nie natywnie', 'Nie', 'Nie', 'Tak (ekosystem Huawei)'],
            ['Victron VRM', 'Nie natywnie', 'Nie', 'Nie', 'Tak (otwarta arch.)'],
            ['MICOMA EMS', 'Nie', 'Nie', 'Nie', 'TAK – polski'],
        ]
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Rekomendacja: ENNO-EMS ')
    run.bold = True
    p.add_run(
        'jako pierwszy wybór (natywna integracja z TGE, polski producent). '
        'Alternatywa: Tibo Energy Alice (najlepsza AI – 45-61% oszczędności vs systemy rule-based).'
    )

    # 4.4 AC vs DC coupling
    doc.add_heading('4.4. Integracja z istniejącym PV: AC-coupled vs DC-coupled', level=2)

    add_styled_table(doc,
        ['Parametr', 'AC-coupled', 'DC-coupled'],
        [
            ['Sprawność', '90-94% (3 konwersje)', '96-98% (1 konwersja)'],
            ['Retrofit do istniejącego PV', 'ŁATWY – brak zmian w PV', 'TRUDNY – wymaga falownika hybrydowego'],
            ['Elastyczność producentów', 'Wysoka', 'Niska (jeden ekosystem)'],
            ['Koszt instalacji', 'Wyższy (dodatkowy falownik)', 'Niższy o 10-20%'],
            ['Serwis', 'Łatwiejszy', 'Zintegrowany'],
        ]
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Dla istniejącego PV: AC-coupled ')
    run.bold = True
    p.add_run('– nie wymaga modyfikacji działającego systemu PV, łatwa instalacja i serwis.')

    # 4.5 Sizing
    doc.add_heading('4.5. Dobór pojemności magazynu', level=2)

    add_styled_table(doc,
        ['Scenariusz', 'Zalecany stosunek BESS/PV'],
        [
            ['Autokonsumpcja PV (podstawowa)', '0,5-1,0 kWh BESS na 1 kWp PV'],
            ['Peak shaving + autokonsumpcja', '1,0-2,0 kWh BESS na 1 kWp PV'],
            ['Pełna optymalizacja + arbitraż', '2,0-4,0 kWh BESS na 1 kWp PV'],
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        'Przykład: zakład z PV 200 kWp, profil 2-zmianowy, nadwyżka PV ~100 kWh/dzień w lecie, '
        'szczyt wieczorny 150 kW × 3h = 450 kWh → rekomendowany magazyn: 250-500 kWh '
        '(z korektą na degradację: × 1,25).'
    )

    doc.add_page_break()

    # ===== 5. ANALIZA RYNKOWA =====
    doc.add_heading('5. Analiza rynkowa', level=1)

    # 5.1 Ceny TGE
    doc.add_heading('5.1. Ceny energii na TGE', level=2)

    add_styled_table(doc,
        ['Parametr', 'Wartość'],
        [
            ['Średnia RDN 2024', '424,94 PLN/MWh'],
            ['Średnia RDN 2023', '533,62 PLN/MWh'],
            ['Średnia RDN listopad 2025', '568,46 PLN/MWh'],
            ['Indeks TGePVm (profil PV, maj 2024)', '265,6 PLN/MWh'],
            ['Taryfa URE na 2026', '495,16 PLN/MWh'],
            ['Kontrakty terminowe Q1 2026', '~405 PLN/MWh'],
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        'Profil dobowy: najniższe ceny 10:00-15:00 (efekt „duck curve" – generacja PV), '
        'najwyższe 17:00-21:00 (wieczorny szczyt). W weekendy profil jest bardziej płaski z jeszcze niższymi cenami w południe.'
    )

    # 5.2 Ujemne ceny
    doc.add_heading('5.2. Ujemne ceny energii w Polsce', level=2)

    add_styled_table(doc,
        ['Rok', 'Godziny z ujemnymi cenami', 'Trend'],
        [
            ['2023', '~30 godzin', 'Początek zjawiska'],
            ['2024', '~186 godzin', '6-krotny wzrost'],
            ['I poł. 2025', '251 godzin (więcej niż cały 2024)', 'Eksplozja'],
            ['2025 (szacunki)', '>315 godzin', 'Dalszy wzrost'],
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        'Rekord: -37 000 PLN/MWh (30 lipca 2025, godz. 10:00-10:15). Typowe ujemne ceny na RDN: '
        '-50 do -200 PLN/MWh. Zjawisko występuje głównie w godzinach 10:00-15:00, w weekendy i święta, '
        'wiosną i latem. Trend jednoznacznie rosnący – przy 25,5 GW PV i 11,2 GW wiatru w Polsce.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Znaczenie dla BESS: ')
    run.bold = True
    p.add_run(
        'Magazyn energii może ładować się za darmo lub nawet ZARABIAĆ na ładowaniu w godzinach ujemnych cen, '
        'a następnie rozładowywać w szczycie po 500-800 PLN/MWh.'
    )

    # 5.3 Arbitraż
    doc.add_heading('5.3. Arbitraż cenowy', level=2)

    add_styled_table(doc,
        ['Scenariusz', 'Spread dzienny', 'Przychód netto/dzień (2 MWh, RTE 85%)', 'Przychód roczny'],
        [
            ['Konserwatywny', '200 PLN/MWh', '340 PLN', '~124 000 PLN'],
            ['Umiarkowany', '350 PLN/MWh', '595 PLN', '~217 000 PLN'],
            ['Optymistyczny (z ujemnymi cenami)', '600 PLN/MWh', '1 020 PLN', '~372 000 PLN'],
        ]
    )

    # 5.4 Peak shaving
    doc.add_heading('5.4. Peak shaving i opłata mocowa', level=2)

    doc.add_paragraph('Stawki opłaty mocowej:')
    add_styled_table(doc,
        ['Rok', 'Stawka (PLN/MWh)', 'Zmiana r/r'],
        [
            ['2024', '~120', '–'],
            ['2025', '141,20', '+17%'],
            ['2026', '219,40', '+55%'],
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        'Opłata mocowa naliczana jest w dni robocze, godz. 7:00-21:59. Odbiorcy przypisywani są do kategorii K1-K4. '
        'Wygładzenie profilu (z K4 na K1) może zmniejszyć opłatę mocową nawet o 83%.'
    )

    add_styled_table(doc,
        ['Kategoria', 'Różnica profilu', 'Płaci % stawki bazowej'],
        [
            ['K1', '< 5%', '17%'],
            ['K2', '5-10%', '40%'],
            ['K3', '10-15%', '70%'],
            ['K4', '≥ 15%', '100%'],
        ]
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Przykład: ')
    run.bold = True
    p.add_run(
        'Firma zużywająca 1 GWh/rok z nierównomiernym profilem (K4) płaci ~160 000 PLN opłaty mocowej. '
        'Po instalacji magazynu i wygładzeniu profilu (K1): ~27 000 PLN. Oszczędność: do 133 000 PLN/rok.'
    )

    # 5.5 Prognozy
    doc.add_heading('5.5. Prognozy cen energii 2025-2027', level=2)

    doc.add_paragraph(
        'Zjawisko „duck curve" będzie się pogłębiać – ceny w południe spadają (rosnący PV), '
        'ceny wieczorne rosną (brak generacji PV + wzrost zapotrzebowania). Spread intraday będzie się '
        'zwiększać, co jest korzystne dla magazynów energii.'
    )

    bullets = [
        'Wzrost: ceny uprawnień CO₂, wyłączanie bloków węglowych, rosnące koszty bilansowania OZE',
        'Spadek średniej: dalszy przyrost OZE, budowa 3,2 GW nowych mocy gazowych',
        'PSE ostrzega: możliwe niedobory mocy do 50 godzin rocznie do 2027 r.',
        'Prognoza 2030: 772 PLN/MWh (przemysł) / 1 124 PLN/MWh (gospodarstwa domowe)',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_page_break()

    # ===== 6. BUSINESS CASE =====
    doc.add_heading('6. Business case – strumienie przychodów', level=1)

    doc.add_paragraph('Zestawienie potencjalnych strumieni przychodów/oszczędności dla magazynu 1 MW / 2 MWh:')

    add_styled_table(doc,
        ['Strumień', 'Roczny przychód (szacunek)', 'Uwagi'],
        [
            ['Arbitraż cenowy (RDN/RDB)', '120 000 – 370 000 PLN', 'Zależy od strategii tradingowej i spreadu'],
            ['Peak shaving (opłata mocowa)', '50 000 – 133 000 PLN', 'Rośnie +55% w 2026 r.'],
            ['Autokonsumpcja PV (nadwyżki)', '30 000 – 100 000 PLN', 'Zależy od wielkości PV i nadwyżek'],
            ['Usługi pomocnicze (aFRR, mFRR)', 'Do 4 000 000 PLN', 'Ryzyko nasycenia rynku'],
            ['Rynek mocy', '~62 000 PLN', 'Po współczynniku KWD 13,39%'],
            ['ŁĄCZNIE (bez usł. pomoc.)', '200 000 – 603 000 PLN/rok', 'Bezpieczny scenariusz'],
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        'Przy CAPEX 3-4 mln PLN za 1 MW/2 MWh, okres zwrotu z arbitrażu + peak shaving + autokonsumpcji: '
        '5-8 lat (bez dotacji). Z dofinansowaniem (np. 45%): 3-5 lat.'
    )

    doc.add_page_break()

    # ===== 7. REKOMENDACJE =====
    doc.add_heading('7. Rekomendacje', level=1)

    recs = [
        ('Taryfa', 'Przejście na taryfę dynamiczną (ceny 15-minutowe TGE) u jednego z 6 zobowiązanych sprzedawców. '
         'Rekomendacja: Pstryk/Fortum (0 zł abonamentu) lub Tauron (niski abonament).'),
        ('Magazyn energii (BESS)', 'Technologia LFP, AC-coupled z istniejącym PV. '
         'Rekomendowane systemy: CATL EnerOne (10 000 cykli), Sungrow PowerStack 255CS, Huawei LUNA2000. '
         'Pojemność: 1-2 kWh na 1 kWp zainstalowanego PV (min. z korektą na degradację × 1,25).'),
        ('Software EMS', 'ENNO-EMS (Ennovation Technology) – natywna integracja z TGE, polski producent. '
         'Strategia: ładowanie w godzinach ujemnych/niskich cen, rozładowanie w szczycie, peak shaving.'),
        ('Dotacje', 'Monitorowanie naborów: Energia Plus (możliwe wznowienie), regionalne RPO, FEnIKS. '
         'Ulga termomodernizacyjna dla JDG.'),
        ('Następne kroki', '1) Analiza faktury i profilu zużycia zakładu. '
         '2) Audyt techniczny istniejącej instalacji PV. '
         '3) Symulacja sizing magazynu (HOMER Pro / PVsyst). '
         '4) Zapytanie ofertowe do integratorów (BESST POWER, Solarpro, 7SUN). '
         '5) Złożenie wniosku o warunki przyłączenia do OSD.'),
    ]

    for title, desc in recs:
        doc.add_heading(title, level=3)
        doc.add_paragraph(desc)

    # ===== ŹRÓDŁA =====
    doc.add_page_break()
    doc.add_heading('Źródła', level=1)

    sources = [
        'TGE – Rynek Dnia Następnego: tge.pl/energia-elektryczna-rdn',
        'TGE – Ceny dynamiczne: tge.pl/ceny_dynamiczne',
        'Enea – Ceny Dynamiczne Firma: enea.pl/ceny-dynamiczne-firma',
        'Energa – Oferta Dynamiczna dla Firm: energa.pl/mala-firma/oferty/oferta-dynamiczna-dla-firm',
        'URE – Taryfy na 2026 r.: ure.gov.pl',
        'Gramwzielone – Ujemne ceny energii: gramwzielone.pl',
        'WNP – Ujemne ceny energii: wnp.pl',
        'Axpo – Rekordowe ujemne ceny: axpo.com/pl',
        'PV Force – Peak Shaving: pvforce.pl/blog/peak-shaving',
        'Alians OZE – Opłata mocowa 2026: alians-oze.pl',
        'WysokieNapiecie – Opłata mocowa rośnie o 50%: wysokienapiecie.pl',
        'Dudkowiak – BESS Legal Requirements 2025: dudkowiak.com',
        'CMS Expert Guide – Energy Storage Poland: cms.law',
        'NFOŚiGW – Nabór magazyny energii: gov.pl/web/nfosigw',
        'CATL EnerOne: catl.com',
        'Sungrow PowerStack: sungrow.com',
        'Huawei LUNA2000: solar.huawei.com',
        'Ennovation Technology ENNO-EMS: ennovationtech.eu',
        'Tibo Energy Alice AI: tibo.energy',
        'PSME – Polskie Stowarzyszenie Magazynowania Energii: psme.org.pl',
        'Energy Instrat – Ceny energii: energy.instrat.pl',
        'Clean Horizon – Poland BESS Market: ess-news.com',
        'Capstone DC – Europe Battery Storage: capstonedc.com',
    ]
    for s in sources:
        doc.add_paragraph(s, style='List Bullet')

    # Zapisz
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Raport_PV_BESS_Zaklady_Produkcyjne.docx')
    doc.save(output_path)
    print(f'Raport zapisany: {output_path}')
    return output_path


def create_report_bytes() -> bytes:
    """Generuje raport DOCX i zwraca jako bytes (do st.download_button)."""
    import io
    import tempfile
    path = create_report()
    with open(path, 'rb') as f:
        return f.read()


if __name__ == '__main__':
    create_report()
